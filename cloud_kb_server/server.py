#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BBCode 云知识库服务器（简化版）
无需 API 认证，适合本地或内网使用
支持 PDF、Word 导入转换为 Markdown
"""

import json
import os
import hashlib
import re
from datetime import datetime
from pathlib import Path
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename

app = Flask(__name__)
CORS(app)  # 允许跨域请求

# 配置
DATA_DIR = Path(__file__).parent / "data"
KNOWLEDGE_FILE = DATA_DIR / "knowledge.json"
UPLOADS_DIR = DATA_DIR / "uploads"

# 确保数据目录存在
DATA_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)

# 允许的文件类型
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'md', 'txt', 'json'}


def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def load_knowledge():
    """加载知识库数据"""
    if KNOWLEDGE_FILE.exists():
        with open(KNOWLEDGE_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def save_knowledge(data):
    """保存知识库数据"""
    with open(KNOWLEDGE_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def generate_id(title):
    """生成唯一ID"""
    import re
    # 移除特殊字符，转换为小写，空格替换为连字符
    id_str = re.sub(r'[^\w\s-]', '', title.lower())
    id_str = re.sub(r'[-\s]+', '-', id_str).strip('-')
    
    # 如果为空，使用时间戳
    if not id_str:
        id_str = f"item-{int(datetime.now().timestamp())}"
    
    # 确保唯一性
    knowledge = load_knowledge()
    base_id = id_str[:50]
    final_id = base_id
    counter = 1
    while final_id in knowledge:
        final_id = f"{base_id}-{counter}"
        counter += 1
    
    return final_id


def calculate_checksum(title, content, version):
    """计算校验和"""
    text = f"{title}{content}{version}"
    return hashlib.md5(text.encode('utf-8')).hexdigest()


def pdf_to_markdown(file_path):
    """将 PDF 转换为 Markdown"""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        markdown_content = []
        
        # 提取标题（使用文件名或第一页内容）
        title = Path(file_path).stem
        
        markdown_content.append(f"# {title}\n")
        markdown_content.append(f"\n> 从 PDF 导入 | 共 {len(doc)} 页\n")
        markdown_content.append("\n---\n")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # 提取文本
            text = page.get_text()
            
            if text.strip():
                markdown_content.append(f"\n## 第 {page_num + 1} 页\n")
                
                # 简单格式化：检测标题和列表
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # 检测可能的标题（短行、全大写或以数字开头）
                    if len(line) < 100 and (line.isupper() or re.match(r'^\d+\.', line)):
                        markdown_content.append(f"\n### {line}\n")
                    # 检测列表项
                    elif line.startswith('•') or line.startswith('-') or re.match(r'^\d+\.', line):
                        markdown_content.append(f"{line}\n")
                    else:
                        markdown_content.append(f"{line}\n")
                
                markdown_content.append("\n")
        
        doc.close()
        return title, '\n'.join(markdown_content)
        
    except Exception as e:
        print(f"PDF 转换错误: {e}")
        return None, None


def word_to_markdown(file_path):
    """将 Word 文档转换为 Markdown"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        markdown_content = []
        
        # 提取标题
        title = Path(file_path).stem
        
        markdown_content.append(f"# {title}\n")
        markdown_content.append(f"\n> 从 Word 导入\n")
        markdown_content.append("\n---\n")
        
        # 遍历段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 根据样式判断标题级别
            style_name = para.style.name.lower() if para.style else ''
            
            if 'heading 1' in style_name or para.style.name == '标题 1':
                markdown_content.append(f"\n# {text}\n")
            elif 'heading 2' in style_name or para.style.name == '标题 2':
                markdown_content.append(f"\n## {text}\n")
            elif 'heading 3' in style_name or para.style.name == '标题 3':
                markdown_content.append(f"\n### {text}\n")
            elif text.startswith('•') or text.startswith('-'):
                markdown_content.append(f"- {text[1:].strip()}\n")
            elif re.match(r'^\d+\.', text):
                markdown_content.append(f"{text}\n")
            else:
                markdown_content.append(f"{text}\n")
        
        # 处理表格
        for table in doc.tables:
            markdown_content.append("\n")
            for i, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                
                markdown_content.append("| " + " | ".join(row_text) + " |\n")
                
                # 添加表头分隔符
                if i == 0:
                    markdown_content.append("| " + " | ".join(["---"] * len(row_text)) + " |\n")
            
            markdown_content.append("\n")
        
        return title, '\n'.join(markdown_content)
        
    except Exception as e:
        print(f"Word 转换错误: {e}")
        return None, None


def txt_to_markdown(file_path):
    """将文本文件转换为 Markdown"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        title = Path(file_path).stem
        
        markdown = f"# {title}\n\n"
        markdown += f"> 从文本文件导入\n\n---\n\n"
        markdown += content
        
        return title, markdown
        
    except Exception as e:
        print(f"文本文件转换错误: {e}")
        return None, None


# ============ API 路由 ============

@app.route('/api/v1/health', methods=['GET'])
def health_check():
    """健康检查"""
    return jsonify({
        "status": "ok",
        "version": "1.0.0-simple",
        "timestamp": datetime.now().isoformat(),
        "auth": "disabled"
    })


@app.route('/api/v1/info', methods=['GET'])
def get_info():
    """获取服务器信息"""
    knowledge = load_knowledge()
    categories = set(item['category'] for item in knowledge.values())
    
    return jsonify({
        "name": "BBCode Cloud KB Server (Simple)",
        "version": "1.0.0-simple",
        "total_items": len(knowledge),
        "categories": sorted(list(categories)),
        "auth_required": False
    })


@app.route('/api/v1/knowledge', methods=['GET'])
def list_knowledge():
    """获取知识条目列表"""
    knowledge = load_knowledge()
    items = list(knowledge.values())
    
    # 分类筛选
    category = request.args.get('category')
    if category:
        items = [item for item in items if item['category'] == category]
    
    # 标签筛选
    tags = request.args.get('tags')
    if tags:
        tag_list = [t.strip() for t in tags.split(',')]
        items = [item for item in items if any(tag in item['tags'] for tag in tag_list)]
    
    # 搜索
    search = request.args.get('search')
    if search:
        search_lower = search.lower()
        items = [item for item in items 
                if search_lower in item['title'].lower() 
                or search_lower in item['content'].lower()
                or any(search_lower in tag.lower() for tag in item['tags'])]
    
    # 按更新时间排序
    items.sort(key=lambda x: x['updated_at'], reverse=True)
    
    return jsonify(items)


@app.route('/api/v1/knowledge/<item_id>', methods=['GET'])
def get_knowledge(item_id):
    """获取单个知识条目"""
    knowledge = load_knowledge()
    item = knowledge.get(item_id)
    
    if not item:
        return jsonify({"error": True, "message": "知识条目不存在"}), 404
    
    return jsonify(item)


@app.route('/api/v1/knowledge', methods=['POST'])
def create_knowledge():
    """创建知识条目"""
    data = request.json
    
    if not data or 'title' not in data or 'content' not in data:
        return jsonify({"error": True, "message": "缺少必要字段: title, content"}), 400
    
    item_id = generate_id(data['title'])
    now = datetime.now().isoformat()
    
    item = {
        "id": item_id,
        "title": data['title'],
        "content": data['content'],
        "category": data.get('category', 'general'),
        "tags": data.get('tags', []),
        "version": 1,
        "created_at": now,
        "updated_at": now,
        "checksum": calculate_checksum(data['title'], data['content'], 1),
        "author": data.get('author', 'anonymous')
    }
    
    knowledge = load_knowledge()
    knowledge[item_id] = item
    save_knowledge(knowledge)
    
    return jsonify(item), 201


@app.route('/api/v1/import', methods=['POST'])
def import_document():
    """导入文档并转换为 Markdown"""
    if 'file' not in request.files:
        return jsonify({"error": True, "message": "没有上传文件"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": True, "message": "未选择文件"}), 400
    
    if not allowed_file(file.filename):
        return jsonify({"error": True, "message": "不支持的文件类型"}), 400
    
    # 保存上传的文件
    filename = secure_filename(file.filename)
    file_path = UPLOADS_DIR / filename
    file.save(file_path)
    
    try:
        # 根据文件类型转换
        ext = filename.rsplit('.', 1)[1].lower()
        title = None
        content = None
        
        if ext == 'pdf':
            title, content = pdf_to_markdown(str(file_path))
        elif ext in ['docx', 'doc']:
            title, content = word_to_markdown(str(file_path))
        elif ext == 'txt':
            title, content = txt_to_markdown(str(file_path))
        elif ext == 'md':
            title, content = txt_to_markdown(str(file_path))
        else:
            return jsonify({"error": True, "message": f"暂不支持的文件类型: {ext}"}), 400
        
        if title is None or content is None:
            return jsonify({"error": True, "message": "文件转换失败"}), 500
        
        # 创建知识条目
        item_id = generate_id(title)
        now = datetime.now().isoformat()
        
        item = {
            "id": item_id,
            "title": title,
            "content": content,
            "category": request.form.get('category', 'imported'),
            "tags": ['imported', ext],
            "version": 1,
            "created_at": now,
            "updated_at": now,
            "checksum": calculate_checksum(title, content, 1),
            "author": 'imported'
        }
        
        knowledge = load_knowledge()
        knowledge[item_id] = item
        save_knowledge(knowledge)
        
        # 删除临时文件
        os.remove(file_path)
        
        return jsonify({
            "success": True,
            "message": f"成功导入 {ext.upper()} 文件",
            "item": item
        }), 201
        
    except Exception as e:
        # 清理临时文件
        if file_path.exists():
            os.remove(file_path)
        return jsonify({"error": True, "message": f"导入失败: {str(e)}"}), 500


@app.route('/api/v1/knowledge/<item_id>', methods=['PUT'])
def update_knowledge(item_id):
    """更新知识条目"""
    knowledge = load_knowledge()
    
    if item_id not in knowledge:
        return jsonify({"error": True, "message": "知识条目不存在"}), 404
    
    data = request.json
    item = knowledge[item_id]
    
    # 更新字段
    if 'title' in data:
        item['title'] = data['title']
    if 'content' in data:
        item['content'] = data['content']
    if 'category' in data:
        item['category'] = data['category']
    if 'tags' in data:
        item['tags'] = data['tags']
    if 'author' in data:
        item['author'] = data['author']
    
    item['version'] += 1
    item['updated_at'] = datetime.now().isoformat()
    item['checksum'] = calculate_checksum(item['title'], item['content'], item['version'])
    
    save_knowledge(knowledge)
    
    return jsonify(item)


@app.route('/api/v1/knowledge/<item_id>', methods=['DELETE'])
def delete_knowledge(item_id):
    """删除知识条目"""
    knowledge = load_knowledge()
    
    if item_id not in knowledge:
        return jsonify({"error": True, "message": "知识条目不存在"}), 404
    
    del knowledge[item_id]
    save_knowledge(knowledge)
    
    return jsonify({"message": "删除成功", "id": item_id})


@app.route('/api/v1/categories', methods=['GET'])
def get_categories():
    """获取所有分类"""
    knowledge = load_knowledge()
    categories = set(item['category'] for item in knowledge.values())
    return jsonify(sorted(list(categories)))


@app.route('/api/v1/tags', methods=['GET'])
def get_tags():
    """获取所有标签"""
    knowledge = load_knowledge()
    tags = set()
    for item in knowledge.values():
        tags.update(item['tags'])
    return jsonify(sorted(list(tags)))


# ============ 管理界面 ============

@app.route('/view/<item_id>')
def view_knowledge(item_id):
    """查看知识条目详情"""
    knowledge = load_knowledge()
    item = knowledge.get(item_id)
    
    if not item:
        return "知识条目不存在", 404
    
    html = f'''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{item['title']} - BBCode Cloud KB</title>
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/github-markdown-css@5/github-markdown-dark.min.css">
    <script src="https://cdn.jsdelivr.net/npm/marked@9/marked.min.js"></script>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: #0d1117;
            color: #c9d1d9;
            line-height: 1.6;
            margin: 0;
            padding: 0;
        }}
        .header {{
            background: #161b22;
            border-bottom: 1px solid #30363d;
            padding: 16px 24px;
            position: sticky;
            top: 0;
            z-index: 100;
        }}
        .header-content {{
            max-width: 1200px;
            margin: 0 auto;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 20px;
            color: #58a6ff;
        }}
        .header .meta {{
            color: #8b949e;
            font-size: 14px;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            padding: 32px 24px;
        }}
        .article-header {{
            margin-bottom: 32px;
            padding-bottom: 24px;
            border-bottom: 1px solid #30363d;
        }}
        .article-title {{
            font-size: 32px;
            font-weight: 600;
            color: #f0f6fc;
            margin: 0 0 16px 0;
            line-height: 1.3;
        }}
        .article-meta {{
            display: flex;
            gap: 16px;
            flex-wrap: wrap;
            font-size: 14px;
            color: #8b949e;
        }}
        .article-meta span {{
            display: flex;
            align-items: center;
            gap: 6px;
        }}
        .tag {{
            background: #21262d;
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 12px;
            color: #58a6ff;
        }}
        .markdown-body {{
            background: transparent !important;
            color: #c9d1d9 !important;
        }}
        .markdown-body h1, .markdown-body h2, .markdown-body h3 {{
            color: #f0f6fc;
            border-bottom-color: #30363d;
        }}
        .markdown-body code {{
            background: #161b22;
            color: #ff7b72;
        }}
        .markdown-body pre {{
            background: #161b22;
            border-radius: 8px;
        }}
        .markdown-body blockquote {{
            border-left-color: #3fb950;
            color: #8b949e;
        }}
        .back-btn {{
            background: #21262d;
            color: #c9d1d9;
            border: 1px solid #30363d;
            padding: 8px 16px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 14px;
            text-decoration: none;
            display: inline-flex;
            align-items: center;
            gap: 6px;
        }}
        .back-btn:hover {{
            background: #30363d;
        }}
        .edit-btn {{
            background: #238636;
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 14px;
            text-decoration: none;
            display: inline-flex;
            align-items: center;
            gap: 6px;
        }}
        .edit-btn:hover {{
            background: #2ea043;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="header-content">
            <div>
                <h1>☁️ BBCode Cloud KB</h1>
                <span class="meta">知识库详情</span>
            </div>
            <div style="display: flex; gap: 8px;">
                <a href="/" class="back-btn">← 返回列表</a>
                <button class="edit-btn" onclick="editItem('{item['id']}')">✏️ 编辑</button>
            </div>
        </div>
    </div>
    
    <div class="container">
        <article class="article-header">
            <h1 class="article-title">{item['title']}</h1>
            <div class="article-meta">
                <span>📁 {item['category']}</span>
                <span>🏷️ {' '.join([f'<span class="tag">{t}</span>' for t in item['tags']])}</span>
                <span>🕐 {item['updated_at'][:19].replace('T', ' ')}</span>
                <span>✏️ v{item['version']}</span>
            </div>
        </article>
        
        <div class="markdown-body" id="content"></div>
    </div>
    
    <script>
        const markdownContent = {json.dumps(item['content'], ensure_ascii=False)};
        document.getElementById('content').innerHTML = marked.parse(markdownContent);
        
        function editItem(id) {{
            window.location.href = '/?edit=' + id;
        }}
    </script>
</body>
</html>
    '''
    return html


@app.route('/')
def admin_page():
    """管理界面"""
    html = '''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>BBCode 云知识库管理</title>
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/github-markdown-css@5/github-markdown-dark.min.css">
    <script src="https://cdn.jsdelivr.net/npm/marked@9/marked.min.js"></script>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: #1e1e1e;
            color: #d4d4d4;
            line-height: 1.6;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }
        h1 {
            color: #4ec9b0;
            margin-bottom: 10px;
        }
        .subtitle {
            color: #858585;
            margin-bottom: 30px;
        }
        .stats {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
            margin-bottom: 30px;
        }
        .stat-card {
            background: #252526;
            padding: 20px;
            border-radius: 8px;
            border: 1px solid #333;
        }
        .stat-card h3 {
            color: #858585;
            font-size: 14px;
            margin-bottom: 10px;
        }
        .stat-card .value {
            font-size: 32px;
            color: #4ec9b0;
            font-weight: bold;
        }
        .actions {
            margin-bottom: 20px;
        }
        .btn {
            background: #007acc;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 4px;
            cursor: pointer;
            font-size: 14px;
            margin-right: 10px;
            margin-bottom: 10px;
        }
        .btn:hover {
            background: #0098ff;
        }
        .btn-success {
            background: #4ec9b0;
        }
        .btn-success:hover {
            background: #5dd9c0;
        }
        .btn-danger {
            background: #f44336;
        }
        .btn-danger:hover {
            background: #ff6659;
        }
        .search-box {
            width: 100%;
            padding: 12px;
            background: #3c3c3c;
            border: 1px solid #555;
            border-radius: 4px;
            color: #d4d4d4;
            font-size: 14px;
            margin-bottom: 20px;
        }
        .knowledge-list {
            background: #252526;
            border-radius: 8px;
            border: 1px solid #333;
            overflow: hidden;
        }
        .knowledge-item {
            padding: 15px 20px;
            border-bottom: 1px solid #333;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .knowledge-item:last-child {
            border-bottom: none;
        }
        .knowledge-item:hover {
            background: #2a2d2e;
        }
        .knowledge-info h4 {
            color: #dcdcaa;
            margin-bottom: 5px;
        }
        .knowledge-meta {
            font-size: 12px;
            color: #858585;
        }
        .knowledge-meta span {
            margin-right: 15px;
        }
        .tag {
            display: inline-block;
            background: #3c3c3c;
            padding: 2px 8px;
            border-radius: 3px;
            font-size: 11px;
            margin-right: 5px;
        }
        .knowledge-actions {
            display: flex;
            gap: 10px;
        }
        .btn-small {
            padding: 5px 12px;
            font-size: 12px;
        }
        .modal {
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0,0,0,0.8);
            z-index: 1000;
            justify-content: center;
            align-items: center;
        }
        .modal.active {
            display: flex;
        }
        .modal-content {
            background: #252526;
            padding: 30px;
            border-radius: 8px;
            width: 95%;
            max-width: 1200px;
            max-height: 95vh;
            overflow-y: auto;
        }
        .form-group {
            margin-bottom: 20px;
        }
        .form-group label {
            display: block;
            margin-bottom: 8px;
            color: #d4d4d4;
            font-weight: 500;
        }
        .form-group input,
        .form-group select {
            width: 100%;
            padding: 10px;
            background: #3c3c3c;
            border: 1px solid #555;
            border-radius: 4px;
            color: #d4d4d4;
            font-size: 14px;
        }
        .editor-container {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            height: 500px;
        }
        .editor-pane {
            display: flex;
            flex-direction: column;
        }
        .editor-pane h4 {
            margin: 0 0 10px 0;
            color: #858585;
            font-size: 14px;
        }
        .editor-textarea {
            flex: 1;
            padding: 15px;
            background: #1e1e1e;
            border: 1px solid #555;
            border-radius: 4px;
            color: #d4d4d4;
            font-family: 'Consolas', 'Monaco', monospace;
            font-size: 14px;
            line-height: 1.6;
            resize: none;
            tab-size: 4;
        }
        .preview-pane {
            background: #0d1117;
            border: 1px solid #30363d;
            border-radius: 4px;
            padding: 20px;
            overflow-y: auto;
        }
        .preview-pane .markdown-body {
            background: transparent !important;
        }
        .form-actions {
            display: flex;
            justify-content: flex-end;
            gap: 10px;
            margin-top: 20px;
            padding-top: 20px;
            border-top: 1px solid #333;
        }
        .toolbar {
            display: flex;
            gap: 8px;
            margin-bottom: 10px;
            padding: 8px;
            background: #3c3c3c;
            border-radius: 4px;
        }
        .toolbar button {
            background: #252526;
            border: 1px solid #555;
            color: #d4d4d4;
            padding: 5px 10px;
            border-radius: 3px;
            cursor: pointer;
            font-size: 13px;
        }
        .toolbar button:hover {
            background: #3c3c3c;
        }
        .empty-state {
            text-align: center;
            padding: 60px 20px;
            color: #858585;
        }
        .empty-state h3 {
            margin-bottom: 10px;
        }
        .toast {
            position: fixed;
            bottom: 20px;
            right: 20px;
            padding: 15px 20px;
            border-radius: 4px;
            color: white;
            font-size: 14px;
            z-index: 2000;
            animation: slideIn 0.3s ease;
        }
        .toast.success {
            background: #4ec9b0;
        }
        .toast.error {
            background: #f44336;
        }
        @keyframes slideIn {
            from {
                transform: translateX(100%);
                opacity: 0;
            }
            to {
                transform: translateX(0);
                opacity: 1;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>☁️ BBCode 云知识库管理</h1>
        <p class="subtitle">无需认证，本地/内网部署</p>
        
        <div class="stats">
            <div class="stat-card">
                <h3>知识条目</h3>
                <div class="value" id="totalItems">0</div>
            </div>
            <div class="stat-card">
                <h3>分类</h3>
                <div class="value" id="totalCategories">0</div>
            </div>
            <div class="stat-card">
                <h3>标签</h3>
                <div class="value" id="totalTags">0</div>
            </div>
        </div>
        
        <div class="actions">
            <button class="btn btn-success" onclick="openModal()">+ 新建知识条目</button>
            <button class="btn" onclick="loadData()">🔄 刷新</button>
            <button class="btn" onclick="exportAll()">📥 导出全部</button>
            <button class="btn" onclick="document.getElementById('importFile').click()">📤 导入</button>
            <input type="file" id="importFile" accept=".json,.md,.pdf,.docx,.doc,.txt" style="display: none;" onchange="importFile(this)">
            <button class="btn btn-primary" onclick="document.getElementById('docImportFile').click()">📄 导入文档</button>
            <input type="file" id="docImportFile" accept=".pdf,.docx,.doc,.txt" style="display: none;" onchange="importDocument(this)">
        </div>
        
        <input type="text" class="search-box" id="searchBox" placeholder="搜索知识条目..." onkeyup="searchItems()">
        
        <div class="knowledge-list" id="knowledgeList">
            <div class="empty-state">
                <h3>暂无知识条目</h3>
                <p>点击"新建知识条目"按钮添加内容</p>
            </div>
        </div>
    </div>
    
    <!-- 编辑模态框 -->
    <div class="modal" id="editModal">
        <div class="modal-content">
            <h2 id="modalTitle">新建知识条目</h2>
            <form id="knowledgeForm">
                <input type="hidden" id="itemId">
                <div class="form-group">
                    <label>标题</label>
                    <input type="text" id="itemTitle" required placeholder="输入知识条目标题">
                </div>
                <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 15px;">
                    <div class="form-group">
                        <label>分类</label>
                        <select id="itemCategory">
                            <option value="general">通用</option>
                            <option value="python">Python</option>
                            <option value="javascript">JavaScript</option>
                            <option value="tutorial">教程</option>
                            <option value="tips">技巧</option>
                            <option value="reference">参考</option>
                            <option value="faq">FAQ</option>
                        </select>
                    </div>
                    <div class="form-group">
                        <label>标签（用逗号分隔）</label>
                        <input type="text" id="itemTags" placeholder="例如: python, 基础, 入门">
                    </div>
                </div>
                <div class="form-group">
                    <label>内容（支持 Markdown）</label>
                    <div class="toolbar">
                        <button type="button" onclick="insertMarkdown('**', '**')" title="粗体"><b>B</b></button>
                        <button type="button" onclick="insertMarkdown('*', '*')" title="斜体"><i>I</i></button>
                        <button type="button" onclick="insertMarkdown('# ', '')" title="标题">H1</button>
                        <button type="button" onclick="insertMarkdown('## ', '')" title="标题">H2</button>
                        <button type="button" onclick="insertMarkdown('### ', '')" title="标题">H3</button>
                        <button type="button" onclick="insertMarkdown('- ', '')" title="列表">• List</button>
                        <button type="button" onclick="insertMarkdown('1. ', '')" title="有序列表">1. List</button>
                        <button type="button" onclick="insertMarkdown('```\n', '\n```')" title="代码块">{ }</button>
                        <button type="button" onclick="insertMarkdown('> ', '')" title="引用">" Quote</button>
                        <button type="button" onclick="insertMarkdown('[', '](url)')" title="链接">Link</button>
                        <button type="button" onclick="insertMarkdown('![alt](', ')')" title="图片">Img</button>
                        <button type="button" onclick="insertMarkdown('---\n', '')" title="分隔线">—</button>
                    </div>
                    <div class="editor-container">
                        <div class="editor-pane">
                            <h4>📝 Markdown 编辑器</h4>
                            <textarea id="itemContent" class="editor-textarea" required placeholder="输入 Markdown 内容..." oninput="updatePreview()"></textarea>
                        </div>
                        <div class="editor-pane">
                            <h4>👁️ 实时预览</h4>
                            <div class="preview-pane markdown-body" id="preview"></div>
                        </div>
                    </div>
                </div>
                <div class="form-actions">
                    <button type="button" class="btn" onclick="closeModal()">取消</button>
                    <button type="button" class="btn" onclick="insertTemplate()">📄 插入模板</button>
                    <button type="submit" class="btn btn-success">💾 保存</button>
                </div>
            </form>
        </div>
    </div>
    
    <script>
        let allItems = [];
        
        // 加载数据
        async function loadData() {
            try {
                const response = await fetch('/api/v1/knowledge');
                allItems = await response.json();
                renderItems(allItems);
                updateStats();
            } catch (error) {
                showToast('加载失败: ' + error.message, 'error');
            }
        }
        
        // 更新统计
        async function updateStats() {
            try {
                const [catRes, tagRes] = await Promise.all([
                    fetch('/api/v1/categories'),
                    fetch('/api/v1/tags')
                ]);
                const categories = await catRes.json();
                const tags = await tagRes.json();
                
                document.getElementById('totalItems').textContent = allItems.length;
                document.getElementById('totalCategories').textContent = categories.length;
                document.getElementById('totalTags').textContent = tags.length;
            } catch (error) {
                console.error('更新统计失败:', error);
            }
        }
        
        // 渲染列表
        function renderItems(items) {
            const container = document.getElementById('knowledgeList');
            
            if (items.length === 0) {
                container.innerHTML = `
                    <div class="empty-state">
                        <h3>暂无知识条目</h3>
                        <p>点击"新建知识条目"按钮添加内容</p>
                    </div>
                `;
                return;
            }
            
            container.innerHTML = items.map(item => `
                <div class="knowledge-item">
                    <div class="knowledge-info">
                        <h4>${escapeHtml(item.title)}</h4>
                        <div class="knowledge-meta">
                            <span>📁 ${item.category}</span>
                            <span>🏷️ ${item.tags.map(t => `<span class="tag">${escapeHtml(t)}</span>`).join('')}</span>
                            <span>🕐 ${new Date(item.updated_at).toLocaleString()}</span>
                            <span>v${item.version}</span>
                        </div>
                    </div>
                    <div class="knowledge-actions">
                        <button class="btn btn-small" onclick="viewItem('${item.id}')">👁️ 查看</button>
                        <button class="btn btn-small" onclick="editItem('${item.id}')">✏️ 编辑</button>
                        <button class="btn btn-small btn-danger" onclick="deleteItem('${item.id}')">🗑️ 删除</button>
                    </div>
                </div>
            `).join('');
        }
        
        // 搜索
        function searchItems() {
            const query = document.getElementById('searchBox').value.toLowerCase();
            const filtered = allItems.filter(item => 
                item.title.toLowerCase().includes(query) ||
                item.content.toLowerCase().includes(query) ||
                item.tags.some(t => t.toLowerCase().includes(query))
            );
            renderItems(filtered);
        }
        
        // 打开模态框
        function openModal(item = null) {
            document.getElementById('modalTitle').textContent = item ? '编辑知识条目' : '新建知识条目';
            document.getElementById('itemId').value = item ? item.id : '';
            document.getElementById('itemTitle').value = item ? item.title : '';
            document.getElementById('itemCategory').value = item ? item.category : 'general';
            document.getElementById('itemTags').value = item ? item.tags.join(', ') : '';
            document.getElementById('itemContent').value = item ? item.content : '';
            document.getElementById('editModal').classList.add('active');
            updatePreview();
        }
        
        // 关闭模态框
        function closeModal() {
            document.getElementById('editModal').classList.remove('active');
            document.getElementById('knowledgeForm').reset();
            document.getElementById('preview').innerHTML = '';
        }
        
        // 查看条目
        function viewItem(id) {
            window.open('/view/' + id, '_blank');
        }
        
        // 编辑条目
        function editItem(id) {
            const item = allItems.find(i => i.id === id);
            if (item) openModal(item);
        }
        
        // 更新预览
        function updatePreview() {
            const content = document.getElementById('itemContent').value;
            document.getElementById('preview').innerHTML = marked.parse(content || '### 预览区域\n\n开始输入 Markdown 内容，这里将实时显示预览。');
        }
        
        // 插入 Markdown
        function insertMarkdown(before, after) {
            const textarea = document.getElementById('itemContent');
            const start = textarea.selectionStart;
            const end = textarea.selectionEnd;
            const text = textarea.value;
            const selected = text.substring(start, end);
            
            const newText = text.substring(0, start) + before + selected + after + text.substring(end);
            textarea.value = newText;
            
            textarea.focus();
            textarea.selectionStart = start + before.length;
            textarea.selectionEnd = start + before.length + selected.length;
            
            updatePreview();
        }
        
        // 插入模板
        function insertTemplate() {
            const template = `# 标题

## 简介

在这里输入简介内容...

## 详细内容

### 小节 1

- 要点 1
- 要点 2
- 要点 3

### 小节 2

\`\`\`python
# 代码示例
print("Hello, World!")
\`\`\`

## 参考

- [链接标题](https://example.com)

---

**关键词**: 标签1, 标签2, 标签3`;
            
            document.getElementById('itemContent').value = template;
            updatePreview();
        }
        
        // 检查 URL 参数，如果有 edit 参数则打开编辑
        function checkUrlParams() {
            const params = new URLSearchParams(window.location.search);
            const editId = params.get('edit');
            if (editId) {
                const item = allItems.find(i => i.id === editId);
                if (item) {
                    openModal(item);
                }
            }
        }
        
        // 删除条目
        async function deleteItem(id) {
            if (!confirm('确定要删除这个知识条目吗？')) return;
            
            try {
                const response = await fetch(`/api/v1/knowledge/${id}`, { method: 'DELETE' });
                if (response.ok) {
                    showToast('删除成功');
                    loadData();
                } else {
                    throw new Error('删除失败');
                }
            } catch (error) {
                showToast('删除失败: ' + error.message, 'error');
            }
        }
        
        // 保存表单
        document.getElementById('knowledgeForm').onsubmit = async function(e) {
            e.preventDefault();
            
            const id = document.getElementById('itemId').value;
            const data = {
                title: document.getElementById('itemTitle').value,
                category: document.getElementById('itemCategory').value,
                tags: document.getElementById('itemTags').value.split(',').map(t => t.trim()).filter(t => t),
                content: document.getElementById('itemContent').value
            };
            
            try {
                const url = id ? `/api/v1/knowledge/${id}` : '/api/v1/knowledge';
                const method = id ? 'PUT' : 'POST';
                
                const response = await fetch(url, {
                    method: method,
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify(data)
                });
                
                if (response.ok) {
                    showToast(id ? '更新成功' : '创建成功');
                    closeModal();
                    loadData();
                } else {
                    throw new Error('保存失败');
                }
            } catch (error) {
                showToast('保存失败: ' + error.message, 'error');
            }
        };
        
        // 显示提示
        function showToast(message, type = 'success') {
            const toast = document.createElement('div');
            toast.className = `toast ${type}`;
            toast.textContent = message;
            document.body.appendChild(toast);
            setTimeout(() => toast.remove(), 3000);
        }
        
        // HTML转义
        function escapeHtml(text) {
            const div = document.createElement('div');
            div.textContent = text;
            return div.innerHTML;
        }
        
        // 导出全部
        async function exportAll() {
            try {
                const data = {
                    export_time: new Date().toISOString(),
                    total_items: allItems.length,
                    items: allItems
                };
                
                const blob = new Blob([JSON.stringify(data, null, 2)], { type: 'application/json' });
                const url = URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `knowledge_backup_${new Date().toISOString().slice(0, 10)}.json`;
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                URL.revokeObjectURL(url);
                
                showToast('导出成功');
            } catch (error) {
                showToast('导出失败: ' + error.message, 'error');
            }
        }
        
        // 导入文件
        async function importFile(input) {
            const file = input.files[0];
            if (!file) return;
            
            try {
                const text = await file.text();
                let imported = 0;
                
                if (file.name.endsWith('.json')) {
                    // JSON 格式
                    const data = JSON.parse(text);
                    const items = data.items || data;
                    
                    for (const item of items) {
                        if (item.title && item.content) {
                            await fetch('/api/v1/knowledge', {
                                method: 'POST',
                                headers: { 'Content-Type': 'application/json' },
                                body: JSON.stringify({
                                    title: item.title,
                                    content: item.content,
                                    category: item.category || 'general',
                                    tags: item.tags || [],
                                    author: item.author || 'imported'
                                })
                            });
                            imported++;
                        }
                    }
                } else if (file.name.endsWith('.md')) {
                    // Markdown 格式
                    const title = file.name.replace('.md', '');
                    await fetch('/api/v1/knowledge', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({
                            title: title,
                            content: text,
                            category: 'general',
                            tags: ['imported'],
                            author: 'imported'
                        })
                    });
                    imported = 1;
                }
                
                showToast(`成功导入 ${imported} 条知识`);
                loadData();
            } catch (error) {
                showToast('导入失败: ' + error.message, 'error');
            }
            
            input.value = '';
        }
        
        // 导入文档（PDF、Word、TXT）
        async function importDocument(input) {
            const file = input.files[0];
            if (!file) return;
            
            const allowedTypes = ['.pdf', '.docx', '.doc', '.txt'];
            const ext = '.' + file.name.split('.').pop().toLowerCase();
            
            if (!allowedTypes.includes(ext)) {
                showToast('不支持的文件类型，请上传 PDF、Word 或 TXT 文件', 'error');
                input.value = '';
                return;
            }
            
            const formData = new FormData();
            formData.append('file', file);
            formData.append('category', 'imported');
            
            showToast('正在转换文档，请稍候...');
            
            try {
                const response = await fetch('/api/v1/import', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                if (response.ok) {
                    showToast(result.message || '导入成功');
                    loadData();
                } else {
                    showToast(result.message || '导入失败', 'error');
                }
            } catch (error) {
                showToast('导入失败: ' + error.message, 'error');
            }
            
            input.value = '';
        }
        
        // 点击模态框外部关闭
        document.getElementById('editModal').onclick = function(e) {
            if (e.target === this) closeModal();
        };
        
        // 初始加载
        loadData();
        checkUrlParams();
    </script>
</body>
</html>
    '''
    return html


# 添加示例数据
def add_sample_data():
    """添加示例数据"""
    if KNOWLEDGE_FILE.exists():
        return
    
    sample_data = {
        "python-basics": {
            "id": "python-basics",
            "title": "Python 基础知识",
            "content": """# Python 基础

Python 是一种高级编程语言，具有简洁易读的语法。

## 变量定义

```python
name = "Python"
age = 30
is_awesome = True
```

## 数据类型

- int: 整数
- float: 浮点数
- str: 字符串
- bool: 布尔值
- list: 列表
- dict: 字典""",
            "category": "python",
            "tags": ["python", "基础", "入门"],
            "version": 1,
            "created_at": "2024-01-01T00:00:00",
            "updated_at": "2024-01-01T00:00:00",
            "checksum": "abc123",
            "author": "system"
        },
        "code-style": {
            "id": "code-style",
            "title": "代码风格指南",
            "content": """# 代码风格指南

## 命名规范

- 变量名：小写字母，下划线分隔
- 函数名：小写字母，下划线分隔
- 类名：驼峰命名法

## 注释规范

- 行内注释：# 后加一个空格
- 文档字符串：使用三引号

## 代码格式化

- 每行不超过 79 个字符
- 使用 4 个空格缩进""",
            "category": "tips",
            "tags": ["规范", "风格", "最佳实践"],
            "version": 1,
            "created_at": "2024-01-01T00:00:00",
            "updated_at": "2024-01-01T00:00:00",
            "checksum": "def456",
            "author": "system"
        }
    }
    
    save_knowledge(sample_data)
    print("✅ 已添加示例数据")


def get_local_ip():
    """获取本机局域网 IP 地址"""
    import socket
    try:
        # 创建一个 UDP 套接字
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        # 连接到一个公共 DNS 服务器
        s.connect(("8.8.8.8", 80))
        # 获取本地 IP
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"


if __name__ == '__main__':
    import argparse
    
    # 命令行参数解析
    parser = argparse.ArgumentParser(description='BBCode 云知识库服务器')
    parser.add_argument('--host', default='0.0.0.0', help='服务器主机地址 (默认: 0.0.0.0)')
    parser.add_argument('--port', type=int, default=5000, help='服务器端口 (默认: 5000)')
    parser.add_argument('--no-debug', action='store_true', help='关闭调试模式')
    args = parser.parse_args()
    
    local_ip = get_local_ip()
    
    print("🚀 启动 BBCode 云知识库服务器（简化版）")
    print("=" * 60)
    print(f"📁 数据目录: {DATA_DIR}")
    print("🔓 认证: 已禁用（无需 API Key）")
    print("-" * 60)
    print("🌐 访问地址:")
    print(f"   本机访问: http://127.0.0.1:{args.port}")
    print(f"   本机访问: http://localhost:{args.port}")
    if local_ip != "127.0.0.1":
        print(f"   局域网访问: http://{local_ip}:{args.port}")
    print("=" * 60)
    print("💡 提示: 按 Ctrl+C 停止服务器")
    print("=" * 60)
    
    # 添加示例数据
    add_sample_data()
    
    # 启动服务器
    app.run(
        host=args.host, 
        port=args.port, 
        debug=not args.no_debug,
        threaded=True  # 启用多线程，支持并发访问
    )
